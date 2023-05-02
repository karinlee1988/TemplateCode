#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Time : 2022/2/8 9:44
# @Author : karinlee
# @FileName : excel套打至word.py
# @Software : PyCharm
# @Blog : https://blog.csdn.net/weixin_43972976

"""
用于套打用。将一个excel数据表格多行的内容分别套打到word文档中，excel数据表格每行生成一份套打的word文档

"""
import os
import time
import openpyxl
from docx import Document

class ExcelPrintToWord(object):
    """
    20230502 test OK
    """

    def __init__(self, word_path, excel_path, isprint=False):
        """

        :param word_path:  作为套打模板的word文档全文件名
        :type word_path: str

        :param excel_path: 作为数据来源的excel文档全文件名
        :type excel_path: str

        :param isprint: 是否需要实际操作打印机执行打印命令，默认为否
        :type isprint: bool

        """

        self.word_path = word_path
        self.excel_path = excel_path
        self.isprint = isprint
        # 获取excel文档工作薄、工作表对象
        self.xlsx = openpyxl.load_workbook(self.excel_path)
        self.xlsx_sheet = self.xlsx.active

    @staticmethod
    def replace_text(doc: [Document], old_text: str, new_text: str) -> None:
        """
        将某个文档的某个词进行替换

        :param doc:   要处理的文件
        :type doc: class Document

        :param old_text:  旧的词
        :type old_text: str

        :param new_text: 新的词
        :type new_text: str

        :return: None
        """
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in inline:
                    if old_text in i.text:
                        text = i.text.replace(old_text, new_text)
                        i.text = text

    def main(self):
        """
        程序入口，可以根据实际情况修改后进行套打
        """
        # -------------------
        # 这里修改套打数据（excel文档）的开始编号->结束编号
        # 编号也等于序号，用于套打时能方便定位寻找相应文档。一般来说，excel表第一行是表头，从第二行开始为编号（序号）1的数据，
        # 第三行为编号（序号）2的数据...
        tag_begin = 1
        tag_end = 10
        # -------------------
        # -------------------
        # 这里修改套打时要替换的元素内容，这里准备了5个元素，可以根据实际的表格和文档结构增加或删除元素
        for row in range(tag_begin + 1, tag_end + 2):

            tag = str(self.xlsx_sheet.cell(row=row, column=1).value)
            sjr = str(self.xlsx_sheet.cell(row=row, column=2).value)
            phone = str(self.xlsx_sheet.cell(row=row, column=3).value)
            address = str(self.xlsx_sheet.cell(row=row, column=4).value)
            company = str(self.xlsx_sheet.cell(row=row, column=5).value)

            # 每次生成一个新的class Document()
            docx = Document(self.word_path)
            self.replace_text(docx, "{tag}", tag)
            self.replace_text(docx, "{sjr}", sjr)
            self.replace_text(docx, "{address}", address)
            self.replace_text(docx, "{phone}", phone)
            self.replace_text(docx, "{company}", company)
            docx.save(tag + "_通知.docx")
        # -------------------
            # 是否需要实际打印，默认为否
            if self.isprint:
                os.startfile(tag + "_通知.docx", 'print')
                # 根据打印机的情况设置延时
                time.sleep(10)

if __name__ == '__main__':
    p = ExcelPrintToWord('tests\\通知.docx', 'tests\\名单.xlsx', isprint=False)
    p.main()